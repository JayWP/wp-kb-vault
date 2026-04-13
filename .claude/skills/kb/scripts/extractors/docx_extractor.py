"""Extract text from Word documents."""
from docx import Document
import os


def extract(docx_path: str, output_dir: str) -> str:
    """Extract all paragraphs and tables from docx."""
    basename = os.path.splitext(os.path.basename(docx_path))[0]
    txt_path = os.path.join(output_dir, f"{basename}.txt")

    doc = Document(docx_path)
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for i, table in enumerate(doc.tables):
        parts.append(f"\n--- Table {i+1} ---")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("\n".join(parts))

    print(f"  Word: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables extracted")
    return txt_path
